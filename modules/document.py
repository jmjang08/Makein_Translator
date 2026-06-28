# modules/document.py
from io import BytesIO
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.text import WD_BREAK
from PIL import Image


class Paragraph:
    def __init__(self, image: Image.Image = None, text: str = "", bold: bool = False, page_break: bool = False):
        self.image: Image.Image = image
        self.text: str = text
        self.bold: bool = bold
        self.page_break: bool = page_break

    def _add_text_with_line_breaks(self, paragraph):
        """줄바꿈 유지"""
        normalized_text = self.text.replace("\r\n", "\n").replace("\r", "\n")
        for idx, line in enumerate(normalized_text.split("\n")):
            if idx:
                paragraph.add_run().add_break(WD_BREAK.LINE)
            run = paragraph.add_run(line if line else "\u00A0")
            run.bold = self.bold
    
    def to_docx_paragraph(self, document: Document):
        """DOCX 문단 추가"""
        paragraph = document.add_paragraph()

        if self.page_break:
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            return paragraph

        if self.text:
            self._add_text_with_line_breaks(paragraph)

        if self.image:
            buffer = BytesIO()
            image_format = self.image.format or "PNG"
            self.image.save(buffer, format=image_format)
            buffer.seek(0)
            section = document.sections[0]
            available_width = section.page_width - section.left_margin - section.right_margin
            available_height = section.page_height - section.top_margin - section.bottom_margin
            emu_per_inch = 914400
            max_width_in = available_width / emu_per_inch
            max_height_in = available_height / emu_per_inch
            dpi_x, dpi_y = self.image.info.get("dpi", (72, 72))
            width_px, height_px = self.image.size
            width_in = width_px / (dpi_x or 72)
            height_in = height_px / (dpi_y or 72)
            scale = min(
                max_width_in / width_in if width_in else 1.0,
                max_height_in / height_in if height_in else 1.0,
                1.0,
            )
            target_width = width_in * scale
            target_height = height_in * scale
            paragraph.add_run().add_picture(
                buffer,
                width=Inches(target_width),
                height=Inches(target_height),
            )

        return paragraph

class Docx:
    def __init__(self):
        self.doc: list[Paragraph] = []

    def _iter_run_elements(self, paragraph):
        """run 요소 순회"""
        w_r = qn('w:r')
        w_hyperlink = qn('w:hyperlink')
        w_fldsimple = qn('w:fldSimple')

        for child in paragraph._p:
            if child.tag == w_r:
                yield child
            elif child.tag in (w_hyperlink, w_fldsimple):
                for run in child.iterchildren():
                    if run.tag == w_r:
                        yield run

    def _run_text(self, run_element) -> str:
        """run 텍스트 추출"""
        texts = []
        for node in run_element.iter():
            if node.tag == qn('w:t'):
                texts.append(node.text or "")
            elif node.tag == qn('w:tab'):
                texts.append("\t")
            elif node.tag in (qn('w:br'), qn('w:cr')):
                if node.tag == qn('w:br') and node.get(qn('w:type')) == 'page':
                    continue
                texts.append("\n")
        return "".join(texts)

    def _run_has_page_break(self, run_element) -> bool:
        """run 페이지 나누기 확인"""
        for node in run_element.iter():
            if node.tag == qn('w:br') and node.get(qn('w:type')) == 'page':
                return True
        return False

    def _paragraph_has_page_break(self, paragraph) -> bool:
        """문단 페이지 나누기 확인"""
        ppr = paragraph._p.find(qn('w:pPr'))
        if ppr is not None and ppr.find(qn('w:pageBreakBefore')) is not None:
            return True
        for node in paragraph._p.iter():
            if node.tag == qn('w:br') and node.get(qn('w:type')) == 'page':
                return True
        return False

    def _run_is_bold(self, run_element) -> bool:
        """굵게 서식 확인"""
        rpr = run_element.find(qn('w:rPr'))
        if rpr is None:
            return False
        bold = rpr.find(qn('w:b'))
        if bold is None:
            return False
        val = bold.get(qn('w:val'))
        if val is None:
            return True
        return str(val).lower() not in ("0", "false", "off")

    def load_from_path(self, file_path: str, max_len: int):
        """DOCX 불러오기"""
        self.doc = []
        document = Document(file_path)

        buffer_text = ""
        buffer_any_bold = False
        buffer_all_bold = True

        def flush_buffer():
            """버퍼 비우기"""
            nonlocal buffer_text, buffer_any_bold, buffer_all_bold
            if buffer_text:
                bold_value = buffer_all_bold and buffer_any_bold
                self.doc.append(Paragraph(text=buffer_text, bold=bold_value))
                buffer_text = ""
                buffer_any_bold = False
                buffer_all_bold = True

        def append_page_break():
            """페이지 나누기 추가"""
            if self.doc and self.doc[-1].page_break:
                return
            self.doc.append(Paragraph(page_break=True))

        def append_segment(text: str, seg_any_bold: bool, seg_all_bold: bool, new_paragraph: bool):
            """텍스트 세그먼트 추가"""
            nonlocal buffer_text, buffer_any_bold, buffer_all_bold
            if not text:
                return
            separator = "\n" if (new_paragraph and buffer_text) else ""
            combined = f"{buffer_text}{separator}{text}" if buffer_text else text
            if max_len <= 0 or len(combined) <= max_len:
                buffer_text = combined
                buffer_any_bold = buffer_any_bold or seg_any_bold
                buffer_all_bold = buffer_all_bold and seg_all_bold
                return

            if buffer_text and max_len > 0 and len(combined) > max_len:
                flush_buffer()
                buffer_text = ""
                buffer_any_bold = False
                buffer_all_bold = True
                combined = text
                if len(combined) <= max_len:
                    buffer_text = combined
                    buffer_any_bold = seg_any_bold
                    buffer_all_bold = seg_all_bold
                    return
            remaining = combined
            while remaining:
                chunk = remaining[:max_len] if max_len > 0 else remaining
                remaining = remaining[len(chunk):]
                buffer_text = chunk
                buffer_any_bold = seg_any_bold
                buffer_all_bold = seg_all_bold
                flush_buffer()

        for paragraph in document.paragraphs:
            text_buffer = []
            any_bold = False
            all_bold = True
            paragraph_has_content = False
            paragraph_has_image = False

            if self._paragraph_has_page_break(paragraph):
                flush_buffer()
                append_page_break()

            is_new_paragraph = bool(buffer_text)

            def flush_text_segment():
                """텍스트 세그먼트 비우기"""
                nonlocal text_buffer, any_bold, all_bold, is_new_paragraph
                if text_buffer:
                    segment_text = "".join(text_buffer)
                    append_segment(segment_text, any_bold, all_bold, is_new_paragraph)
                    text_buffer = []
                    any_bold = False
                    all_bold = True
                    is_new_paragraph = False

            for run_el in self._iter_run_elements(paragraph):
                if self._run_has_page_break(run_el):
                    flush_text_segment()
                    flush_buffer()
                    append_page_break()
                run_text = self._run_text(run_el)
                if run_text:
                    cleaned_text = run_text.replace("\x0c", "")
                    if cleaned_text:
                        text_buffer.append(cleaned_text)
                        paragraph_has_content = True
                        is_bold = self._run_is_bold(run_el)
                        any_bold = any_bold or is_bold
                        all_bold = all_bold and is_bold

                has_pic = any(
                    node.tag.endswith('}pic')
                    for node in run_el.iter()
                )
                if has_pic:
                    flush_text_segment()
                    flush_buffer()
                    paragraph_has_image = True
                    blips = [
                        node for node in run_el.iter()
                        if node.tag.endswith('}blip')
                    ]
                    for blip in blips:
                        embed = blip.get(qn('r:embed'))
                        if not embed:
                            continue
                        image_part = document.part.related_parts.get(embed)
                        if not image_part:
                            continue
                        image_bytes = image_part.blob
                        image = Image.open(BytesIO(image_bytes))
                        self.doc.append(Paragraph(image=image))
                    is_new_paragraph = False

            flush_text_segment()
            if not paragraph_has_content and not paragraph_has_image:
                append_segment("\n", False, True, False)
                continue

        flush_buffer()
        return self.doc
    
    def save_to_path(self, file_path: str):
        """DOCX 저장"""
        document = Document()
        section = document.sections[0]
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        for paragraph in self.doc:
            paragraph.to_docx_paragraph(document)
        document.save(file_path)


if __name__ == "__main__":
    docx = Docx()
    docx.load_from_path('./target/test.docx')
    docx.save_to_path('./target/test_out.docx')
